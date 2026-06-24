from docx import Document
import os

os.makedirs('reports/phase1', exist_ok=True)

# 1. Phase_1_Report.docx
doc1 = Document()
doc1.add_heading('FakeShield AI: Phase 1 Report', 0)

doc1.add_heading('Project Title', level=1)
doc1.add_paragraph('FakeShield AI: Privacy-Preserving Fake News Detection using Federated Learning and BERT')

doc1.add_heading('Problem Statement', level=1)
doc1.add_paragraph('Traditional fake news detection systems require centralized collection of news data, which raises significant privacy concerns and risks data breaches. Organizations are hesitant to share their proprietary data. Therefore, there is a need for a decentralized approach that can train robust AI models across multiple organizations without exposing raw data.')

doc1.add_heading('Objectives', level=1)
doc1.add_paragraph('1. Develop a highly accurate Fake News Detection model using DistilBERT.\n2. Implement a Federated Learning framework (Flower) to train the model across multiple distributed clients.\n3. Ensure data privacy by keeping raw text data locally on the client machines.\n4. Evaluate the performance of the global model against traditional centralized approaches.')

doc1.add_heading('Technologies Used', level=1)
doc1.add_paragraph('- Programming Language: Python 3.10+\n- Deep Learning Framework: PyTorch\n- NLP Model: HuggingFace Transformers (DistilBERT)\n- Federated Learning: Flower (flwr)\n- Data Processing: Pandas, NumPy\n- Visualization: Matplotlib, Seaborn')

doc1.add_heading('Dataset Details', level=1)
doc1.add_paragraph('The system utilizes the ISOT Fake News Dataset, containing thousands of articles categorized into Real and Fake news. The dataset is partitioned among multiple clients to simulate a real-world federated learning scenario.')

doc1.add_heading('Literature Review Summary', level=1)
doc1.add_paragraph('Extensive research has shown that transformer-based models like BERT achieve state-of-the-art results in NLP classification tasks. Simultaneously, Federated Learning has emerged as the leading paradigm for privacy-preserving machine learning. Combining these two approaches allows for high-accuracy fake news detection while maintaining strict data confidentiality.')

doc1.add_heading('Proposed Methodology', level=1)
doc1.add_paragraph('The methodology involves initializing a global DistilBERT model on a central server. Multiple clients download this model, train it locally on their private datasets, and send only the updated model weights back to the server. The server aggregates these weights using the Federated Averaging (FedAvg) algorithm to create an improved global model.')

doc1.add_heading('Current Progress', level=1)
doc1.add_paragraph('As of Phase 1, the project topic has been finalized. The dataset has been collected and preprocessed. The initial Federated Learning code and DistilBERT model have been successfully executed and tested in a simulated environment (Google Colab). The codebase is maintained on GitHub.')

doc1.save('reports/phase1/Phase_1_Report.docx')

# 2. Dataset_Information.docx
doc2 = Document()
doc2.add_heading('Dataset Information', 0)

doc2.add_heading('Dataset Name', level=1)
doc2.add_paragraph('ISOT Fake News Dataset (Fake.csv and True.csv)')

doc2.add_heading('Source', level=1)
doc2.add_paragraph('University of Victoria ISOT Research Lab / Kaggle')

doc2.add_heading('Number of Records', level=1)
doc2.add_paragraph('Total: ~44,898 articles\n- True News: ~21,417 articles\n- Fake News: ~23,481 articles')

doc2.add_heading('Features', level=1)
doc2.add_paragraph('1. title: The title of the news article.\n2. text: The main body content of the article.\n3. subject: The category of the news (e.g., politics, world news).\n4. date: The date the article was published.')

doc2.add_heading('Labels', level=1)
doc2.add_paragraph('- 1 (True): Represents authentic, real news articles.\n- 0 (Fake): Represents fabricated or deceptive news articles.')

doc2.add_heading('Sample Dataset Table', level=1)
table = doc2.add_table(rows=3, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Title'
hdr_cells[1].text = 'Text Snippet'
hdr_cells[2].text = 'Label'

row_cells = table.rows[1].cells
row_cells[0].text = 'Scientists discover new planet'
row_cells[1].text = 'NASA announced the discovery...'
row_cells[2].text = 'True (1)'

row_cells = table.rows[2].cells
row_cells[0].text = 'Aliens land in New York'
row_cells[1].text = 'Extraterrestrials have arrived...'
row_cells[2].text = 'Fake (0)'

doc2.add_heading('Dataset Screenshots', level=1)
doc2.add_paragraph('[Please insert screenshot of the Pandas DataFrame head() output here]')

doc2.save('reports/phase1/Dataset_Information.docx')

# 3. Literature_Review.docx
doc3 = Document()
doc3.add_heading('Literature Review', 0)

papers = [
    {
        "title": "Fake News Detection using Deep Learning",
        "authors": "Goldani et al.",
        "year": "2021",
        "summary": "This paper proposes CNN and capsule networks for detecting fake news. It focuses on extracting textual features effectively.",
        "advantages": "High accuracy on standard datasets; handles local text context well.",
        "limitations": "Does not address data privacy; requires centralized data collection.",
        "ref": "Goldani, M. H., Safabakhsh, R., & Momtazi, S. (2021). Convolutional neural network with margin loss for fake news detection. Information Processing & Management."
    },
    {
        "title": "Federated Learning for Privacy-Preserving NLP",
        "authors": "Lin et al.",
        "year": "2022",
        "summary": "Explores the application of federated learning in various natural language processing tasks to preserve user privacy.",
        "advantages": "Strong privacy guarantees; reduces the need for data transfer.",
        "limitations": "High communication overhead and slower convergence times.",
        "ref": "Lin, B. Y., et al. (2022). FedNLP: Benchmarking Federated Learning Methods for Natural Language Processing Tasks."
    },
    {
        "title": "Fake News Detection using BERT",
        "authors": "Jwa et al.",
        "year": "2019",
        "summary": "Utilizes the pre-trained BERT model to capture complex contextual relationships in news articles to detect misinformation.",
        "advantages": "State-of-the-art accuracy; captures deep semantic meaning.",
        "limitations": "Computationally expensive to train on centralized servers.",
        "ref": "Jwa, H., et al. (2019). exBAKE: Automatic Fake News Detection Model Based on Bidirectional Encoder Representations from Transformers (BERT)."
    },
    {
        "title": "Communication-Efficient Learning of Deep Networks from Decentralized Data",
        "authors": "McMahan et al.",
        "year": "2017",
        "summary": "The foundational paper introducing the Federated Averaging (FedAvg) algorithm for training distributed networks.",
        "advantages": "Pioneered the field of Federated Learning; highly robust algorithm.",
        "limitations": "Struggles with highly Non-IID (unbalanced) data distributions.",
        "ref": "McMahan, B., et al. (2017). Communication-Efficient Learning of Deep Networks from Decentralized Data. AISTATS."
    },
    {
        "title": "A Survey on Federated Learning in Privacy-Preserving Data Mining",
        "authors": "Yang et al.",
        "year": "2019",
        "summary": "Provides a comprehensive overview of federated learning architectures (horizontal, vertical, and federated transfer learning).",
        "advantages": "Categorizes different FL approaches clearly.",
        "limitations": "A survey paper, so it does not propose a specific novel model.",
        "ref": "Yang, Q., et al. (2019). Federated Machine Learning: Concept and Applications. ACM Transactions on Intelligent Systems and Technology."
    }
]

for p in papers:
    doc3.add_heading(p['title'], level=2)
    doc3.add_paragraph(f"Authors: {p['authors']}\nYear: {p['year']}")
    doc3.add_paragraph(f"Summary: {p['summary']}")
    doc3.add_paragraph(f"Advantages: {p['advantages']}")
    doc3.add_paragraph(f"Limitations: {p['limitations']}")
    doc3.add_paragraph(f"Reference: {p['ref']}")
    doc3.add_paragraph("")

doc3.save('reports/phase1/Literature_Review.docx')

# 4. Project_Diary.docx
doc4 = Document()
doc4.add_heading('Project Diary (Phase 1)', 0)

doc4.add_paragraph('Week 1: Topic Finalized', style='Heading 2')
doc4.add_paragraph('- Discussed potential AI project ideas with teammates and guide.\n- Finalized "FakeShield AI: Privacy-Preserving Fake News Detection using Federated Learning and BERT" as the final year project.\n- Approved by the project coordinator.', style='List Bullet')

doc4.add_paragraph('Week 2: Literature Review Completed', style='Heading 2')
doc4.add_paragraph('- Read various IEEE and Springer research papers on Fake News detection, BERT, and Federated Learning.\n- Identified the gap: existing models ignore data privacy.\n- Summarized 5 key research papers for the Phase 1 report.', style='List Bullet')

doc4.add_paragraph('Week 3: Dataset Collected', style='Heading 2')
doc4.add_paragraph('- Searched Kaggle and academic repositories for a suitable dataset.\n- Downloaded the ISOT Fake News dataset containing True.csv and Fake.csv.\n- Analyzed the dataset structure and cleaned null values.', style='List Bullet')

doc4.add_paragraph('Week 4: Initial Code Executed & GitHub Repository Created', style='Heading 2')
doc4.add_paragraph('- Set up the project structure in Python.\n- Created the GitHub repository to manage version control.\n- Successfully ran the initial Federated Learning simulation in Google Colab using the Flower framework and DistilBERT.', style='List Bullet')

doc4.add_paragraph('Future Work (Phase 2 & 3)', style='Heading 2')
doc4.add_paragraph('- Train the model on the full dataset with 10 clients.\n- Develop a backend API using FastAPI.\n- Build a frontend dashboard using Streamlit.\n- Generate comparative performance graphs.', style='List Bullet')

doc4.save('reports/phase1/Project_Diary.docx')

# 5. One Page Faculty Review Sheet.docx
doc5 = Document()
doc5.add_heading('One Page Faculty Review Sheet', 0)

doc5.add_paragraph('Project Title:', style='Heading 2')
doc5.add_paragraph('FakeShield AI: Privacy-Preserving Fake News Detection using Federated Learning and BERT')

doc5.add_paragraph('Phase 1 Status:', style='Heading 2')
doc5.add_paragraph('[\u2713] Topic Finalized\n[\u2713] Dataset Collected\n[\u2713] Literature Review Completed\n[\u2713] GitHub Code Executed')

doc5.add_paragraph('Technologies:', style='Heading 2')
doc5.add_paragraph('- Python\n- BERT (DistilBERT)\n- Federated Learning\n- PyTorch\n- Flower Framework')

doc5.add_paragraph('Dataset:', style='Heading 2')
doc5.add_paragraph('ISOT Fake News Dataset (Fake.csv and True.csv)')

doc5.add_paragraph('Next Phase:', style='Heading 2')
doc5.add_paragraph('- Full Experiments & Model Tuning\n- Final Results Generation\n- Comparative Analysis with Centralized Models')

doc5.save('reports/phase1/Faculty_Review_Sheet.docx')

print("Successfully generated all Phase 1 DOCX deliverables.")
